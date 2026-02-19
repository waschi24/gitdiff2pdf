#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
gitdiff2pdf.py - Generate a PR-like PDF (and optionally Word .docx) from unified git diffs.

Highlights:
- Unified view (default): '-' red, '+' green, context gray - Bitbucket-like.
- Optional: Side-by-Side (--view side-by-side).
- Clean baseline layout: file badge, hunk header, text and line numbers align.
- Robust parser: diff --git / --- +++ / rename / hunks; meta lines ignored.
- Hunk header suffix (e.g., function/context after @@) is rendered as a green line
  below the blue header (not inside it).
- Encoding: UTF-8 / UTF-16 (PowerShell '>').
- Fonts: System fonts (Windows: Consolas/Segoe UI; Linux: DejaVu) with safe fallbacks.
- Efficient pagination: multiple files/hunks per page, subtle block gaps, widow/orphan protection.
- Keep-together: per FILE and per HUNK (no splitting if the block fits on a new page).
- Word output: --word / --word-output path.docx  (requires python-docx).
"""

from __future__ import annotations

import argparse
import datetime as dt
import os
import platform
import re
import sys
from dataclasses import dataclass, field
from typing import List, Optional, Tuple

import fitz  # PyMuPDF


# -------------------- Utilities --------------------

def rgb(r: int, g: int, b: int) -> Tuple[float, float, float]:
    return (r / 255.0, g / 255.0, b / 255.0)


def text_width(s: str, fontname: str, fontsize: float) -> float:
    return fitz.get_text_length(s, fontname=fontname, fontsize=fontsize)


# Remove invisible/problematic Unicode globally (BOM / zero-width / NBSP / rare spaces).
_INVIS_GLOBAL = "".join([
    "\ufeff",                    # BOM
    "\u200b\u200c\u200d\u2060",  # zero-widths
    "\u00a0\u202f\u205f",        # NBSP + narrow + medium spaces
    "\u2000\u2001\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200a",  # various thin/em/en spaces
])
INVISIBLE_TRANS = dict.fromkeys(map(ord, _INVIS_GLOBAL), None)

def strip_invisibles(s: str) -> str:
    return s.translate(INVISIBLE_TRANS)

def clean_leading_artifacts(text: str) -> str:
    """
    Remove ellipsis/dot/bullet artifacts only at the very start of the document,
    without touching real code content later.
    Handles: '···', '...', '…', '•', '‧', '∙', '⋅'.
    """
    t = text.lstrip()
    patterns = ["···", "...", "\u2026", "•", "\u2022", "‧", "\u2027", "∙", "\u2219", "⋅", "\u22c5"]
    for p in patterns:
        if t.startswith(p):
            t = t[len(p):].lstrip()
            break
    if len(t) < len(text):
        return t
    return text


def read_text(path: str) -> str:
    """Read file or '-' (STDIN) with robust encoding fallback, then sanitize."""
    data = sys.stdin.buffer.read() if path == "-" else open(path, "rb").read()
    for enc in ("utf-8", "utf-8-sig", "utf-16", "utf-16-le", "utf-16-be", "latin-1"):
        try:
            text = data.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = data.decode("utf-8", errors="replace")
    text = strip_invisibles(text)
    text = clean_leading_artifacts(text)
    return text


def norm_lines(s: str) -> List[str]:
    """Normalize line breaks but keep trailing newlines (for precise line counting)."""
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    return s.splitlines(True)


def sanitize_path(s: str) -> str:
    """Cut at the first non-path-safe character (guards against copy/paste artifacts)."""
    allowed = set("abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789._- /\\")
    out = []
    for ch in s:
        if ch in allowed:
            out.append(ch)
        else:
            break
    return "".join(out).strip()


def safe_font(fontname: str, fallback: str) -> str:
    """Return a PyMuPDF-safe font name (no spaces), or fallback to avoid exceptions."""
    try:
        if (not fontname) or (" " in fontname) or ("\t" in fontname):
            return fallback
        fitz.get_text_length("M", fontname=fontname, fontsize=10)
        return fontname
    except Exception:
        return fallback


# -------------------- Theme --------------------

@dataclass(frozen=True)
class Theme:
    ui_text: Tuple[float, float, float]
    ui_subtle: Tuple[float, float, float]
    header_line: Tuple[float, float, float]

    bg_added: Tuple[float, float, float]
    bg_removed: Tuple[float, float, float]
    bg_context: Tuple[float, float, float]
    bg_hunk: Tuple[float, float, float]

    tx_added: Tuple[float, float, float]
    tx_removed: Tuple[float, float, float]
    tx_context: Tuple[float, float, float]
    tx_header: Tuple[float, float, float]
    tx_hunk: Tuple[float, float, float]

    bar_added: Tuple[float, float, float]
    bar_removed: Tuple[float, float, float]


LIGHT = Theme(
    ui_text=rgb(40, 40, 40),
    ui_subtle=rgb(125, 125, 125),
    header_line=rgb(210, 210, 210),
    bg_added=rgb(225, 245, 234),
    bg_removed=rgb(252, 232, 232),
    bg_context=rgb(247, 247, 249),
    bg_hunk=rgb(232, 240, 252),
    tx_added=rgb(22, 125, 57),
    tx_removed=rgb(178, 36, 30),
    tx_context=rgb(60, 60, 60),
    tx_header=rgb(230, 140, 0),
    tx_hunk=rgb(30, 90, 200),
    bar_added=rgb(34, 170, 84),
    bar_removed=rgb(220, 64, 52),
)

DARK = Theme(
    ui_text=rgb(230, 230, 230),
    ui_subtle=rgb(170, 170, 170),
    header_line=rgb(80, 80, 80),
    bg_added=rgb(36, 64, 52),
    bg_removed=rgb(72, 40, 40),
    bg_context=rgb(36, 36, 40),
    bg_hunk=rgb(44, 60, 84),
    tx_added=rgb(170, 235, 190),
    tx_removed=rgb(255, 170, 170),
    tx_context=rgb(230, 230, 230),
    tx_header=rgb(255, 200, 130),
    tx_hunk=rgb(160, 190, 255),
    bar_added=rgb(60, 200, 110),
    bar_removed=rgb(240, 90, 80),
)


# -------------------- Fonts --------------------

@dataclass
class Fonts:
    ui: str
    ui_bold: str
    mono: str
    mono_bold: str


def load_font_from_file(path: Optional[str]) -> Optional[str]:
    if not path or not os.path.isfile(path):
        return None
    try:
        return fitz.Font(fontfile=path).name
    except Exception:
        return None


def detect_system_fonts() -> Fonts:
    """Windows: Consolas/Segoe UI; Linux: DejaVu; fallback: Courier."""
    ui = "courier"; ui_b = "courier-bold"
    mono = "courier"; mono_b = "courier-bold"

    system = platform.system().lower()
    if "windows" in system:
        wins = os.environ.get("WINDIR", r"C:\Windows")
        cand = {
            "mono": [rf"{wins}\Fonts\consola.ttf", rf"{wins}\Fonts\cour.ttf"],
            "mono_b": [rf"{wins}\Fonts\consolab.ttf", rf"{wins}\Fonts\courbd.ttf"],
            "ui": [rf"{wins}\Fonts\segoeui.ttf", rf"{wins}\Fonts\arial.ttf"],
            "ui_b": [rf"{wins}\Fonts\segoeuib.ttf", rf"{wins}\Fonts\arialbd.ttf"],
        }
        mono = load_font_from_file(next((p for p in cand["mono"] if os.path.isfile(p)), None)) or mono
        mono_b = load_font_from_file(next((p for p in cand["mono_b"] if os.path.isfile(p)), None)) or mono_b
        ui = load_font_from_file(next((p for p in cand["ui"] if os.path.isfile(p)), None)) or ui
        ui_b = load_font_from_file(next((p for p in cand["ui_b"] if os.path.isfile(p)), None)) or ui_b
    elif "linux" in system:
        mono = load_font_from_file("/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf") or mono
        mono_b = load_font_from_file("/usr/share/fonts/truetype/dejavu/DejaVuSansMono-Bold.ttf") or mono_b
        ui = load_font_from_file("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf") or ui
        ui_b = load_font_from_file("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf") or ui_b

    return Fonts(ui=ui, ui_bold=ui_b, mono=mono, mono_bold=mono_b)


# -------------------- Layout --------------------

@dataclass
class Layout:
    margin: float = 44.0
    font_size: float = 9.5
    line_gap: float = 2.2
    hunk_gap_y: float = 8.0
    section_gap_y: float = 10.0
    col_gap: float = 16.0
    gutter_gap: float = 6.0
    gutter_chars: int = 5
    gap_badge_to_hunk: float = 2.0
    gap_hunk_to_code: float = 4.0
    block_gap_y: float = 6.0
    min_rows_on_page: int = 3


# -------------------- Diff Model & Parser --------------------

@dataclass
class DiffLine:
    kind: str       # 'ctx' | 'del' | 'add'
    text: str
    raw: str
    old_num: Optional[int] = None
    new_num: Optional[int] = None


@dataclass
class Hunk:
    header: str
    old_start: int
    old_count: int
    new_start: int
    new_count: int
    lines: List[DiffLine] = field(default_factory=list)
    suffix: Optional[str] = None


@dataclass
class DiffFile:
    old_path: str = ""
    new_path: str = ""
    hunks: List[Hunk] = field(default_factory=list)


HUNK_RE = re.compile(r"@@\s*-(\d+)(?:,(\d+))?\s+\+(\d+)(?:,(\d+))?\s*@@(?:(.*))?$")


def parse_path_from_diff_git(line: str) -> Tuple[Optional[str], Optional[str]]:
    parts = line.strip().split()
    if len(parts) >= 4 and parts[0] == "diff" and parts[1] == "--git":
        a = parts[2][2:] if parts[2].startswith("a/") else parts[2]
        b = parts[3][2:] if parts[3].startswith("b/") else parts[3]
        return a, b
    return None, None


def parse_path_line(line: str) -> Optional[str]:
    rest = line[4:].strip()
    if rest == "/dev/null":
        return None
    if rest.startswith("a/") or rest.startswith("b/"):
        rest = rest[2:]
    return strip_invisibles(sanitize_path(rest)) or None


def parse_unified_diff(text: str, tabsize: int, debug: bool = False) -> List[DiffFile]:
    lines = norm_lines(text)

    files: List[DiffFile] = []
    current: Optional[DiffFile] = None
    current_hunk: Optional[Hunk] = None
    saw_any_hunk = False

    rename_from: Optional[str] = None
    rename_to: Optional[str] = None

    for raw in lines:
        line = strip_invisibles(raw.rstrip("\n"))

        if (line.startswith("index ")
            or line.startswith("new file mode")
            or line.startswith("deleted file mode")
            or line.startswith("Binary files ")):
            continue
        if line.startswith("\\ No newline at end of file"):
            continue

        if line.startswith("diff --git "):
            if current:
                files.append(current)
            a, b = parse_path_from_diff_git(line)
            current = DiffFile(old_path=a or "", new_path=b or "")
            current_hunk = None
            rename_from = None
            rename_to = None
            continue

        if line.startswith("rename from "):
            rename_from = line[len("rename from "):].strip()
            continue
        if line.startswith("rename to "):
            rename_to = line[len("rename to "):].strip()
            continue

        if line.startswith("--- "):
            if current is None:
                current = DiffFile()
            p = parse_path_line(line)
            if p is not None:
                current.old_path = p
            continue

        if line.startswith("+++ "):
            if current is None:
                current = DiffFile()
            p = parse_path_line(line)
            if p is not None:
                current.new_path = p
            if rename_from and not current.old_path:
                current.old_path = rename_from
            if rename_to and not current.new_path:
                current.new_path = rename_to
            continue

        m = HUNK_RE.match(line)
        if m:
            saw_any_hunk = True
            if current is None:
                current = DiffFile()
            old_start = int(m.group(1)); old_count = int(m.group(2) or "1")
            new_start = int(m.group(3)); new_count = int(m.group(4) or "1")
            suffix = (m.group(5) or "").strip()

            header_core = "@@ -{}".format(old_start)
            if m.group(2):
                header_core += f",{old_count}"
            header_core += " +{}".format(new_start)
            if m.group(4):
                header_core += f",{new_count}"
            header_core += " @@"

            current_hunk = Hunk(
                header=header_core,
                old_start=old_start,
                old_count=old_count,
                new_start=new_start,
                new_count=new_count,
                suffix=suffix or None,
            )
            current.hunks.append(current_hunk)
            continue

        if current_hunk is None:
            continue

        if line.startswith("+") and not line.startswith("+++ "):
            current_hunk.lines.append(DiffLine(kind="add", text=strip_invisibles(line[1:].expandtabs(tabsize)), raw=line))
        elif line.startswith("-") and not line.startswith("--- "):
            current_hunk.lines.append(DiffLine(kind="del", text=strip_invisibles(line[1:].expandtabs(tabsize)), raw=line))
        else:
            t = line[1:] if line.startswith(" ") else line
            current_hunk.lines.append(DiffLine(kind="ctx", text=strip_invisibles(t.expandtabs(tabsize)), raw=line))

    if current:
        files.append(current)

    if not saw_any_hunk:
        return []

    for f in files:
        if not f.old_path and not f.new_path:
            f.old_path = "(Unnamed OLD)"
            f.new_path = "(Unnamed NEW)"

    for f in files:
        for h in f.hunks:
            old_ln = h.old_start
            new_ln = h.new_start
            for dl in h.lines:
                if dl.kind == "ctx":
                    dl.old_num = old_ln; dl.new_num = new_ln
                    old_ln += 1; new_ln += 1
                elif dl.kind == "del":
                    dl.old_num = old_ln; old_ln += 1
                elif dl.kind == "add":
                    dl.new_num = new_ln; new_ln += 1

    return files


# -------------------- Wrapping --------------------

def wrap_text(s: str, max_w: float, fontname: str, fontsize: float) -> List[str]:
    if not s:
        return [""]
    if text_width(s, fontname, fontsize) <= max_w:
        return [s]
    out: List[str] = []
    rest = s
    while rest:
        if text_width(rest, fontname, fontsize) <= max_w:
            out.append(rest); break
        lo, hi = 1, len(rest); cut = 1
        while lo <= hi:
            mid = (lo + hi) // 2
            if text_width(rest[:mid], fontname, fontsize) <= max_w:
                cut = mid; lo = mid + 1
            else:
                hi = mid - 1
        slice_ = rest[:cut]
        ws = max(slice_.rfind(" "), slice_.rfind("\t"))
        if ws >= 0 and ws >= int(0.6 * cut):
            out.append(slice_[:ws].rstrip()); rest = rest[ws + 1 :]
        else:
            out.append(slice_); rest = rest[cut:]
    return out


# -------------------- PDF Renderer --------------------

class Renderer:
    def __init__(self, theme: Theme, layout: Layout, landscape: bool, fonts: Fonts):
        self.theme = theme
        self.layout = layout
        self.landscape = landscape
        self.doc = fitz.open()

        self.page: Optional[fitz.Page] = None
        self.y_base: Optional[float] = None
        self.title: Optional[str] = None

        self.ui_font   = safe_font(fonts.ui,       fallback="courier")
        self.ui_bold   = safe_font(fonts.ui_bold,  fallback="courier-bold")
        self.mono_font = safe_font(fonts.mono,     fallback="courier")
        self.mono_bold = safe_font(fonts.mono_bold, fallback="courier-bold")

    def new_page(self) -> fitz.Page:
        base = fitz.paper_rect("a4")
        if self.landscape:
            return self.doc.new_page(width=base.height, height=base.width)
        return self.doc.new_page(width=base.width, height=base.height)

    def box(self, page: fitz.Page) -> Tuple[float, float, float, float]:
        m = self.layout.margin
        return (m, m, page.rect.width - m, page.rect.height - m)

    def page_capacity(self) -> float:
        assert self.page is not None
        x0, y0, x1, y1 = self.box(self.page)
        return y1 - (y0 + 28)

    def draw_header(self, page: fitz.Page, title: str):
        x0, y0, x1, _ = self.box(page)
        fs = self.layout.font_size
        page.insert_text((x0, y0), title, fontname=self.ui_bold, fontsize=fs + 2, color=self.theme.ui_text)
        stamp = dt.datetime.now().strftime("%Y-%m-%d %H:%M")
        tw = text_width(stamp, self.ui_font, fs)
        page.insert_text((x1 - tw, y0), stamp, fontname=self.ui_font, fontsize=fs, color=self.theme.ui_subtle)
        y = y0 + fs + 2 + fs + 4
        page.draw_line((x0, y - 4), (x1, y - 4), color=self.theme.header_line, width=0.8)

    def start_if_needed(self, title: str):
        if self.page is None:
            self.page = self.new_page()
            self.title = title
            self.draw_header(self.page, title)
            x0, y0, x1, y1 = self.box(self.page)
            fs = self.layout.font_size
            self.y_base = y0 + 28 + fs

    def draw_footer_page_numbers(self):
        total = self.doc.page_count
        if total == 0:
            return
        fs = self.layout.font_size
        for i in range(total):
            page = self.doc.load_page(i)
            x0, _, x1, y1 = self.box(page)
            label = f"Page {i + 1} / {total}"
            tw = text_width(label, self.ui_font, fs)
            page.insert_text(((x0 + x1) / 2 - tw / 2, y1), label, fontname=self.ui_font, fontsize=fs, color=self.theme.ui_subtle)

    def space_left(self) -> float:
        assert self.page is not None and self.y_base is not None
        _, _, _, y1 = self.box(self.page)
        fs = self.layout.font_size
        return y1 - (self.y_base - fs)

    def ensure_y(self, rows_h: float):
        assert self.page is not None and self.y_base is not None and self.title is not None
        if self.space_left() >= rows_h:
            return
        self.page = self.new_page()
        self.draw_header(self.page, self.title)
        _, ny0, _, _ = self.box(self.page)
        fs = self.layout.font_size
        self.y_base = ny0 + 28 + fs

    def widow_check_before_hunk(self, line_h: float):
        needed = line_h + self.layout.gap_hunk_to_code + (self.layout.min_rows_on_page * line_h)
        if self.space_left() < needed:
            self.ensure_y(10_000)

    def draw_file_badge(self, label: str):
        assert self.page is not None and self.y_base is not None
        x0, _, x1, _ = self.box(self.page)
        fs = self.layout.font_size
        line_h = fs + self.layout.line_gap
        pad_x = 7

        top = self.y_base - fs
        bottom = top + line_h
        tw = text_width(label, self.ui_bold, fs)
        rect = fitz.Rect(x0, top, min(x1, x0 + pad_x + tw + 7), bottom)
        self.page.draw_rect(rect, fill=self.theme.bg_hunk, color=None, fill_opacity=0.9)
        self.page.insert_text((x0 + pad_x, self.y_base), label, fontname=self.ui_bold, fontsize=fs, color=self.theme.tx_hunk)
        self.y_base = bottom + self.layout.gap_badge_to_hunk + fs

    def draw_hunk_header(self, header: str):
        assert self.page is not None and self.y_base is not None
        x0, _, x1, _ = self.box(self.page)
        fs = self.layout.font_size
        line_h = fs + self.layout.line_gap
        pad_x = 7

        top = self.y_base - fs
        bottom = top + line_h

        self.page.draw_rect(fitz.Rect(x0, top, x1, bottom), fill=self.theme.bg_hunk, color=None, fill_opacity=0.9)
        self.page.insert_text((x0 + pad_x, self.y_base), header, fontname=self.mono_font, fontsize=fs, color=self.theme.tx_hunk)
        self.y_base = bottom + self.layout.gap_hunk_to_code + fs

    def measure_hunk_height_unified(self, hunk: Hunk, hide_context: bool) -> float:
        assert self.page is not None
        x0, y0, x1, y1 = self.box(self.page)
        fs = self.layout.font_size
        line_h = fs + self.layout.line_gap

        sample = f"{'9' * self.layout.gutter_chars} "
        gutter_w = text_width(sample, self.mono_font, fs)
        max_w = max(12.0, x1 - (x0 + gutter_w + self.layout.gutter_gap))

        total = 0.0
        total += line_h + self.layout.gap_hunk_to_code
        if hunk.suffix:
            parts = wrap_text(hunk.suffix, max_w, self.mono_font, fs)
            total += len(parts) * line_h
        for dl in hunk.lines:
            if hide_context and dl.kind == "ctx":
                continue
            parts = wrap_text(dl.text, max_w, self.mono_font, fs)
            total += len(parts) * line_h
        total += self.layout.block_gap_y
        return total

    def measure_file_height_unified(self, diff_file: DiffFile, hide_context: bool) -> float:
        assert self.page is not None
        x0, y0, x1, y1 = self.box(self.page)
        fs = self.layout.font_size
        line_h = fs + self.layout.line_gap

        sample = f"{'9' * self.layout.gutter_chars} "
        gutter_w = text_width(sample, self.mono_font, fs)
        max_w = max(12.0, x1 - (x0 + gutter_w + self.layout.gutter_gap))

        total = 0.0
        total += line_h + self.layout.gap_badge_to_hunk
        for h in diff_file.hunks:
            total += line_h + self.layout.gap_hunk_to_code
            if h.suffix:
                parts = wrap_text(h.suffix, max_w, self.mono_font, fs)
                total += len(parts) * line_h
            for dl in h.lines:
                if hide_context and dl.kind == "ctx":
                    continue
                parts = wrap_text(dl.text, max_w, self.mono_font, fs)
                total += len(parts) * line_h
            total += self.layout.block_gap_y
        total += self.layout.block_gap_y
        return total

    def render_file_unified(self, diff_file: DiffFile, title: str, hide_context: bool):
        self.start_if_needed(title)
        assert self.page is not None and self.y_base is not None

        required = self.measure_file_height_unified(diff_file, hide_context)
        capacity = self.page_capacity()
        if required <= capacity and required > self.space_left():
            self.ensure_y(10_000)

        x0, y0, x1, y1 = self.box(self.page)
        fs = self.layout.font_size
        line_h = fs + self.layout.line_gap

        sample = f"{'9' * self.layout.gutter_chars} "
        gutter_w = text_width(sample, self.mono_font, fs)
        text_x = x0 + gutter_w + self.layout.gutter_gap

        label = diff_file.new_path or diff_file.old_path or "(Unnamed)"
        self.ensure_y(line_h)
        self.draw_file_badge(label)

        for h in diff_file.hunks:
            h_req = self.measure_hunk_height_unified(h, hide_context)
            fresh_capacity = self.page_capacity()
            if h_req > self.space_left() and h_req <= fresh_capacity:
                self.ensure_y(10_000)

            self.widow_check_before_hunk(line_h)
            self.ensure_y(line_h)
            self.draw_hunk_header(h.header)

            if h.suffix:
                x0, y0, x1, y1 = self.box(self.page)
                sample = f"{'9' * self.layout.gutter_chars} "
                gutter_w = text_width(sample, self.mono_font, fs)
                text_x = x0 + gutter_w + self.layout.gutter_gap
                max_w = max(12.0, x1 - text_x)
                parts = wrap_text(h.suffix, max_w, self.mono_font, fs)
                rows_h = len(parts) * line_h
                self.ensure_y(rows_h)
                x0, y0, x1, y1 = self.box(self.page)
                text_x = x0 + gutter_w + self.layout.gutter_gap
                for part in parts:
                    top = self.y_base - fs
                    bottom = top + line_h
                    self.page.draw_rect(fitz.Rect(x0, top, x1, bottom), fill=self.theme.bg_added, color=None, fill_opacity=0.9)
                    self.page.draw_rect(fitz.Rect(x0, top, x0 + 2.5, bottom), fill=self.theme.bar_added, color=None, fill_opacity=1.0)
                    self.page.insert_text((text_x, self.y_base), part, fontname=self.mono_font, fontsize=fs, color=self.theme.tx_added)
                    self.y_base += line_h

            for dl in h.lines:
                if hide_context and dl.kind == "ctx":
                    continue

                max_w = max(12.0, x1 - text_x)
                parts = wrap_text(dl.text, max_w, self.mono_font, fs)
                rows_h = len(parts) * line_h

                self.ensure_y(rows_h)
                x0, y0, x1, y1 = self.box(self.page)
                text_x = x0 + gutter_w + self.layout.gutter_gap

                for j, part in enumerate(parts):
                    top = self.y_base - fs
                    bottom = top + line_h
                    if dl.kind == "add":
                        self.page.draw_rect(fitz.Rect(x0, top, x1, bottom), fill=self.theme.bg_added, color=None, fill_opacity=0.9)
                        self.page.draw_rect(fitz.Rect(x0, top, x0 + 2.5, bottom), fill=self.theme.bar_added, color=None, fill_opacity=1.0)
                    elif dl.kind == "del":
                        self.page.draw_rect(fitz.Rect(x0, top, x1, bottom), fill=self.theme.bg_removed, color=None, fill_opacity=0.9)
                        self.page.draw_rect(fitz.Rect(x0, top, x0 + 2.5, bottom), fill=self.theme.bar_removed, color=None, fill_opacity=1.0)
                    else:
                        self.page.draw_rect(fitz.Rect(x0, top, x1, bottom), fill=self.theme.bg_context, color=None, fill_opacity=0.9)

                    if j == 0:
                        num = dl.new_num if dl.kind == "add" else (dl.old_num if dl.kind == "del" else dl.old_num)
                        if num is not None:
                            ln_text = f"{num:>{self.layout.gutter_chars}d} "
                            self.page.insert_text((x0, self.y_base), ln_text, fontname=self.mono_font, fontsize=fs, color=self.theme.ui_subtle)

                    col = self.theme.tx_context
                    if dl.kind == "add":
                        col = self.theme.tx_added
                    elif dl.kind == "del":
                        col = self.theme.tx_removed
                    self.page.insert_text((text_x, self.y_base), part, fontname=self.mono_font, fontsize=fs, color=col)

                    self.y_base += line_h

            self.y_base += self.layout.block_gap_y

        self.y_base += self.layout.block_gap_y

    def render_file_sbs(self, diff_file: DiffFile, title: str):
        self.start_if_needed(title)
        assert self.page is not None and self.y_base is not None

        x0, y0, x1, y1 = self.box(self.page)
        fs = self.layout.font_size
        line_h = fs + self.layout.line_gap

        gap = self.layout.col_gap
        col_w = (x1 - x0 - gap) / 2
        sample = f"{'9' * self.layout.gutter_chars} "
        gutter_w = text_width(sample, self.mono_font, fs)
        left_x0 = x0; left_x1 = x0 + col_w
        right_x0 = left_x1 + gap; right_x1 = x1
        left_text_x = left_x0 + gutter_w + self.layout.gutter_gap
        right_text_x = right_x0 + gutter_w + self.layout.gutter_gap

        label = diff_file.new_path or diff_file.old_path or "(Unnamed)"
        self.ensure_y(line_h)
        self.draw_file_badge(label)

        for h in diff_file.hunks:
            self.widow_check_before_hunk(line_h)
            self.ensure_y(line_h)
            self.draw_hunk_header(h.header)

            i = 0
            while i < len(h.lines):
                left: Optional[DiffLine] = None
                right: Optional[DiffLine] = None
                ln = h.lines[i]
                if ln.kind == "del":
                    if i + 1 < len(h.lines) and h.lines[i + 1].kind == "add":
                        left = ln; right = h.lines[i + 1]; i += 2
                    else:
                        left = ln; i += 1
                elif ln.kind == "add":
                    right = ln; i += 1
                else:
                    left = ln; right = ln; i += 1

                l_max = max(12.0, left_x1 - left_text_x)
                r_max = max(12.0, right_x1 - right_text_x)
                l_parts = wrap_text(left.text if left else "", l_max, self.mono_font, fs) if left else [""]
                r_parts = wrap_text(right.text if right else "", r_max, self.mono_font, fs) if right else [""]

                rows = max(len(l_parts), len(r_parts))
                rows_h = rows * line_h
                self.ensure_y(rows_h)

                x0, y0, x1, y1 = self.box(self.page)
                left_x0 = x0; left_x1 = x0 + col_w
                right_x0 = left_x1 + gap; right_x1 = x1
                left_text_x = left_x0 + gutter_w + self.layout.gutter_gap
                right_text_x = right_x0 + gutter_w + self.layout.gutter_gap

                for j in range(rows):
                    top = self.y_base - fs; bottom = top + line_h
                    if left and left.kind == "del":
                        self.page.draw_rect(fitz.Rect(left_x0, top, left_x1, bottom), fill=self.theme.bg_removed, color=None, fill_opacity=0.9)
                        self.page.draw_rect(fitz.Rect(left_x0, top, left_x0 + 2.5, bottom), fill=self.theme.bar_removed, color=None, fill_opacity=1.0)
                    if right and right.kind == "add":
                        self.page.draw_rect(fitz.Rect(right_x0, top, right_x1, bottom), fill=self.theme.bg_added, color=None, fill_opacity=0.9)
                        self.page.draw_rect(fitz.Rect(right_x0, top, right_x0 + 2.5, bottom), fill=self.theme.bar_added, color=None, fill_opacity=1.0)
                    if left and right and left.kind == "ctx" and right.kind == "ctx":
                        self.page.draw_rect(fitz.Rect(left_x0, top, left_x1, bottom), fill=self.theme.bg_context, color=None, fill_opacity=0.9)
                        self.page.draw_rect(fitz.Rect(right_x0, top, right_x1, bottom), fill=self.theme.bg_context, color=None, fill_opacity=0.9)

                    if j == 0 and left and left.kind in ("ctx", "del"):
                        ln_txt = f"{(left.old_num if left.old_num is not None else 0):>{self.layout.gutter_chars}d} "
                        self.page.insert_text((left_x0, self.y_base), ln_txt, fontname=self.mono_font, fontsize=fs, color=self.theme.ui_subtle)
                    if j == 0 and right and right.kind in ("ctx", "add"):
                        rn_txt = f"{(right.new_num if right.new_num is not None else 0):>{self.layout.gutter_chars}d} "
                        self.page.insert_text((right_x0, self.y_base), rn_txt, fontname=self.mono_font, fontsize=fs, color=self.theme.ui_subtle)

                    lp = l_parts[j] if j < len(l_parts) else ""
                    rp = r_parts[j] if j < len(r_parts) else ""
                    if left:
                        col = self.theme.tx_context if left.kind == "ctx" else (self.theme.tx_removed if left.kind == "del" else self.theme.tx_context)
                        self.page.insert_text((left_text_x, self.y_base), lp, fontname=self.mono_font, fontsize=fs, color=col)
                    if right:
                        col = self.theme.tx_context if right.kind == "ctx" else (self.theme.tx_added if right.kind == "add" else self.theme.tx_context)
                        self.page.insert_text((right_text_x, self.y_base), rp, fontname=self.mono_font, fontsize=fs, color=col)
                    self.y_base += line_h

            self.y_base += self.layout.block_gap_y

        self.y_base += self.layout.block_gap_y

    def save(self, output_path: str):
        self.draw_footer_page_numbers()
        self.doc.save(output_path)
        self.doc.close()


# -------------------- Word Renderer --------------------

def _rgb_to_hex(rgb_tuple: Tuple[float, float, float]) -> str:
    """Convert a (0..1, 0..1, 0..1) tuple to a 6-digit hex string (no '#')."""
    r, g, b = rgb_tuple
    return "{:02X}{:02X}{:02X}".format(int(r * 255), int(g * 255), int(b * 255))


def render_word(
    all_files: List[DiffFile],
    title: str,
    output_path: str,
    theme: Theme,
    hide_context: bool,
    view: str,
    font_size_pt: float = 9.5,
):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print(
            "[ERROR] python-docx is not installed. Run: pip install python-docx",
            file=sys.stderr,
        )
        sys.exit(1)

    def hex_color(rgb_tuple: Tuple[float, float, float]) -> str:
        return _rgb_to_hex(rgb_tuple)

    def set_cell_bg(cell, hex_str: str):
        """Set table cell background colour via XML shading."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_str)
        tcPr.append(shd)

    def remove_table_borders(table):
        """Remove all borders from every cell in a table."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        # Table-level border removal
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        tblBorders = tblPr.find(qn("w:tblBorders"))
        if tblBorders is None:
            tblBorders = OxmlElement("w:tblBorders")
            tblPr.append(tblBorders)
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            border.set(qn("w:sz"), "0")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "auto")
            tblBorders.append(border)
        # Cell-level border removal (overrides any cell style)
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.find(qn("w:tcBorders"))
                if tcBorders is None:
                    tcBorders = OxmlElement("w:tcBorders")
                    tcPr.append(tcBorders)
                for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    border = OxmlElement(f"w:{border_name}")
                    border.set(qn("w:val"), "none")
                    border.set(qn("w:sz"), "0")
                    border.set(qn("w:space"), "0")
                    border.set(qn("w:color"), "auto")
                    tcBorders.append(border)

    def add_colored_paragraph(doc, text: str, bg_hex: str, fg_rgb, bold: bool = False, mono: bool = True, indent_pt: float = 0):
        """Add a single-row, single-cell table acting as a highlighted paragraph."""
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        remove_table_borders(table)
        cell = table.cell(0, 0)
        set_cell_bg(cell, bg_hex)

        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        if indent_pt:
            para.paragraph_format.left_indent = Pt(indent_pt)
        run = para.add_run(text)
        run.font.name = "Consolas" if mono else "Calibri"
        run.font.size = Pt(font_size_pt)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(
            int(fg_rgb[0] * 255),
            int(fg_rgb[1] * 255),
            int(fg_rgb[2] * 255),
        )
        return table

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Document title heading
    heading = doc.add_heading(title, level=1)
    heading.runs[0].font.color.rgb = RGBColor(40, 40, 40)
    stamp_para = doc.add_paragraph(dt.datetime.now().strftime("Generated: %Y-%m-%d %H:%M"))
    stamp_para.runs[0].font.size = Pt(font_size_pt - 1)
    stamp_para.runs[0].font.color.rgb = RGBColor(125, 125, 125)
    doc.add_paragraph("")  # spacer

    bg_add_hex = hex_color(theme.bg_added)
    bg_del_hex = hex_color(theme.bg_removed)
    bg_ctx_hex = hex_color(theme.bg_context)
    bg_hunk_hex = hex_color(theme.bg_hunk)

    for diff_file in all_files:
        label = diff_file.new_path or diff_file.old_path or "(Unnamed)"

        # --- File badge ---
        add_colored_paragraph(
            doc, label, bg_hunk_hex, theme.tx_hunk, bold=True, mono=False
        )

        for h in diff_file.hunks:
            # --- Hunk header ---
            add_colored_paragraph(
                doc, h.header, bg_hunk_hex, theme.tx_hunk, bold=False, mono=True
            )

            # --- Hunk suffix (rendered as added line) ---
            if h.suffix:
                add_colored_paragraph(
                    doc, h.suffix, bg_add_hex, theme.tx_added, bold=False, mono=True
                )

            # --- Diff lines ---
            if view == "side-by-side":
                i = 0
                while i < len(h.lines):
                    left: Optional[DiffLine] = None
                    right: Optional[DiffLine] = None
                    ln = h.lines[i]
                    if ln.kind == "del":
                        if i + 1 < len(h.lines) and h.lines[i + 1].kind == "add":
                            left = ln; right = h.lines[i + 1]; i += 2
                        else:
                            left = ln; i += 1
                    elif ln.kind == "add":
                        right = ln; i += 1
                    else:
                        left = ln; right = ln; i += 1

                    tbl = doc.add_table(rows=1, cols=2)
                    tbl.style = "Table Grid"
                    remove_table_borders(tbl)
                    l_cell = tbl.cell(0, 0)
                    r_cell = tbl.cell(0, 1)

                    if left and left.kind == "del":
                        set_cell_bg(l_cell, bg_del_hex)
                    else:
                        set_cell_bg(l_cell, bg_ctx_hex)

                    if right and right.kind == "add":
                        set_cell_bg(r_cell, bg_add_hex)
                    else:
                        set_cell_bg(r_cell, bg_ctx_hex)

                    def _fill_cell(cell, dl: Optional[DiffLine], side: str):
                        para = cell.paragraphs[0]
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)
                        if dl is None:
                            return
                        num = dl.old_num if side == "left" else dl.new_num
                        num_str = f"{num:>5d}  " if num is not None else "        "
                        run_num = para.add_run(num_str)
                        run_num.font.name = "Consolas"
                        run_num.font.size = Pt(font_size_pt)
                        run_num.font.color.rgb = RGBColor(125, 125, 125)

                        fg = (
                            theme.tx_removed if dl.kind == "del"
                            else theme.tx_added if dl.kind == "add"
                            else theme.tx_context
                        )
                        run_txt = para.add_run(dl.text)
                        run_txt.font.name = "Consolas"
                        run_txt.font.size = Pt(font_size_pt)
                        run_txt.font.color.rgb = RGBColor(
                            int(fg[0] * 255), int(fg[1] * 255), int(fg[2] * 255)
                        )

                    _fill_cell(l_cell, left, "left")
                    _fill_cell(r_cell, right, "right")

            else:
                # Unified view
                for dl in h.lines:
                    if hide_context and dl.kind == "ctx":
                        continue
                    bg_hex = bg_add_hex if dl.kind == "add" else (bg_del_hex if dl.kind == "del" else bg_ctx_hex)
                    fg = (
                        theme.tx_added if dl.kind == "add"
                        else theme.tx_removed if dl.kind == "del"
                        else theme.tx_context
                    )
                    num = dl.new_num if dl.kind == "add" else dl.old_num
                    num_str = f"{num:>5d}  " if num is not None else "        "
                    line_text = num_str + dl.text

                    add_colored_paragraph(
                        doc, line_text, bg_hex, fg, bold=False, mono=True
                    )

        doc.add_paragraph("")  # spacer between files

    doc.save(output_path)
    print(f"✓ Word document created: {output_path}")


# -------------------- CLI / Main --------------------

def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        description="PR-like PDF (and optionally Word .docx) from unified git diffs.",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    p.add_argument("inputs", nargs="+", help="Diff file(s) or '-' for STDIN")
    p.add_argument("-o", "--output", required=True, help="Output PDF (e.g., ipa-diff.pdf)")
    p.add_argument("--title", default="Changed Code", help="Document title")
    p.add_argument("--view", choices=["unified", "side-by-side"], default="unified", help="Layout mode")
    p.add_argument("--hide-context", action="store_true", help="Hide context lines (show only +/-) in unified view")
    p.add_argument("--landscape", action="store_true", help="A4 landscape")
    p.add_argument("--font-size", type=float, default=9.5, help="Monospace font size")
    p.add_argument("--tabsize", type=int, default=4, help="Tab width for expansion")
    p.add_argument("--theme", choices=["light", "dark"], default="light", help="Color theme")
    p.add_argument("--debug", action="store_true", help="Parser debug output (stderr)")
    p.add_argument("--word", action="store_true", help="Also generate a Word (.docx) file alongside the PDF")
    p.add_argument("--word-output", default=None, metavar="FILE.docx",
                   help="Custom path for the Word output (default: same base name as --output with .docx)")
    # Optional font overrides
    p.add_argument("--mono-font-file", default=None, help="TTF/OTF monospace (e.g., Consolas)")
    p.add_argument("--mono-bold-font-file", default=None, help="TTF/OTF monospace bold")
    p.add_argument("--ui-font-file", default=None, help="TTF/OTF UI (e.g., Segoe UI / Arial)")
    p.add_argument("--ui-bold-font-file", default=None, help="TTF/OTF UI bold")
    return p.parse_args()


def main():
    args = parse_args()

    theme = LIGHT if args.theme == "light" else DARK

    # System fonts + overrides
    sys_fonts = detect_system_fonts()
    mf = load_font_from_file(args.mono_font_file) or sys_fonts.mono
    mfb = load_font_from_file(args.mono_bold_font_file) or sys_fonts.mono_bold
    uif = load_font_from_file(args.ui_font_file) or sys_fonts.ui
    uifb = load_font_from_file(args.ui_bold_font_file) or sys_fonts.ui_bold
    fonts = Fonts(ui=uif, ui_bold=uifb, mono=mf, mono_bold=mfb)

    layout = Layout(font_size=args.font_size)
    renderer = Renderer(theme=theme, layout=layout, landscape=args.landscape, fonts=fonts)

    # Parse inputs
    all_files: List[DiffFile] = []
    for path in args.inputs:
        try:
            txt = read_text(path)
        except FileNotFoundError:
            print(f"[WARN] File not found: {path}", file=sys.stderr)
            continue
        parsed = parse_unified_diff(txt, tabsize=args.tabsize, debug=args.debug)
        all_files.extend(parsed)

    if not all_files:
        print("[ERROR] No parsable diffs found.", file=sys.stderr)
        print("Hints:", file=sys.stderr)
        print("  • Use a unified diff: e.g., `git diff <commit>` or `git show <commit>`", file=sys.stderr)
        print("  • Not supported: `--word-diff`, `--name-only`, `--name-status`", file=sys.stderr)
        sys.exit(2)

    # Render PDF
    for df in all_files:
        if args.view == "unified":
            renderer.render_file_unified(df, title=args.title, hide_context=args.hide_context)
        else:
            renderer.render_file_sbs(df, title=args.title)

    renderer.save(args.output)

    # Render Word (optional)
    if args.word or args.word_output:
        word_path = args.word_output
        if not word_path:
            base, _ = os.path.splitext(args.output)
            word_path = base + ".docx"
        render_word(
            all_files=all_files,
            title=args.title,
            output_path=word_path,
            theme=theme,
            hide_context=args.hide_context,
            view=args.view,
            font_size_pt=args.font_size,
        )


if __name__ == "__main__":
    main()
